"""
Word Document Reader Utility
Extracts text content from .docx files for easy reference.

Usage:
    python read_docx.py <filename.docx>
    python read_docx.py                     # Lists all .docx files and lets you choose
    python read_docx.py --all               # Reads all .docx files in the folder

Output is saved to a .txt file with the same name.

Dependencies:
    pip install python-docx
"""

import os
import sys
from docx import Document


def extract_text_from_docx(filepath):
    """Extract all text content from a Word document."""
    doc = Document(filepath)
    output = []

    filename = os.path.basename(filepath)
    output.append("=" * 60)
    output.append(f"FILE: {filename}")
    output.append(f"PARAGRAPHS: {len(doc.paragraphs)}")
    output.append(f"TABLES: {len(doc.tables)}")
    output.append("=" * 60)
    output.append("")

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Mark headings
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "")
                output.append(f"[HEADING {level}] {text}")
            else:
                output.append(text)

    # Extract tables
    if doc.tables:
        output.append("")
        output.append("=" * 40)
        output.append("TABLES")
        output.append("=" * 40)

        for i, table in enumerate(doc.tables, 1):
            output.append(f"\n--- Table {i} ---")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                output.append(" | ".join(cells))

    return "\n".join(output)


def list_docx_files(folder):
    """List all .docx files in a folder."""
    files = []
    for f in os.listdir(folder):
        if f.endswith('.docx') and not f.startswith('~$'):
            files.append(f)
    return sorted(files)


def save_output(content, original_filepath):
    """Save extracted content to a text file."""
    base = os.path.splitext(original_filepath)[0]
    output_path = f"{base}_content.txt"

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)

    return output_path


def main():
    folder = os.path.dirname(os.path.abspath(__file__))

    if len(sys.argv) > 1:
        arg = sys.argv[1]

        if arg == '--all':
            files = list_docx_files(folder)
            if not files:
                print("No .docx files found in this folder.")
                return

            print(f"Found {len(files)} Word documents.\n")

            for filename in files:
                filepath = os.path.join(folder, filename)
                print(f"Reading: {filename}...")

                try:
                    content = extract_text_from_docx(filepath)
                    output_path = save_output(content, filepath)
                    print(f"  Saved to: {os.path.basename(output_path)}")
                except Exception as e:
                    print(f"  Error: {e}")

            print("\nDone!")

        elif arg in ('--help', '-h'):
            print(__doc__)

        else:
            # Read specific file
            if not arg.endswith('.docx'):
                arg += '.docx'

            filepath = os.path.join(folder, arg) if not os.path.isabs(arg) else arg

            if not os.path.exists(filepath):
                print(f"File not found: {filepath}")
                return

            print(f"Reading: {arg}...")
            content = extract_text_from_docx(filepath)

            # Print to console
            print("\n" + content)

            # Save to file
            output_path = save_output(content, filepath)
            print(f"\nSaved to: {output_path}")

    else:
        # Interactive mode
        files = list_docx_files(folder)

        if not files:
            print("No .docx files found in this folder.")
            return

        print("Available Word documents:\n")
        for i, f in enumerate(files, 1):
            print(f"  {i}. {f}")

        print(f"\n  A. Read ALL files")
        print(f"  Q. Quit\n")

        choice = input("Enter number or letter: ").strip()

        if choice.upper() == 'Q':
            return
        elif choice.upper() == 'A':
            for filename in files:
                filepath = os.path.join(folder, filename)
                print(f"\nReading: {filename}...")

                try:
                    content = extract_text_from_docx(filepath)
                    output_path = save_output(content, filepath)
                    print(f"  Saved to: {os.path.basename(output_path)}")
                except Exception as e:
                    print(f"  Error: {e}")
        else:
            try:
                idx = int(choice) - 1
                if 0 <= idx < len(files):
                    filename = files[idx]
                    filepath = os.path.join(folder, filename)

                    print(f"\nReading: {filename}...\n")
                    content = extract_text_from_docx(filepath)

                    print(content)

                    output_path = save_output(content, filepath)
                    print(f"\nSaved to: {output_path}")
                else:
                    print("Invalid selection.")
            except ValueError:
                print("Invalid input.")


if __name__ == "__main__":
    main()
